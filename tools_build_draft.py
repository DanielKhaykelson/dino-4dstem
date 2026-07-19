# -*- coding: utf-8 -*-
"""Build the full DINO4DSTEM manuscript draft (draft_v2_full_r3.docx) from the
skeleton + key points, with embedded figures and detailed captions.
Style: matched to Khaykelson et al., Nano Lett. 2025 (no em-dashes; measured,
materials-first; plain method description; careful hedging; clear bottom lines)."""
import os, re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = r"D:/DINOSR/Claude/PaperRun_claude/dino_sr_contrastive"
F = {
    "fig1": ROOT + "/docs/paper/draft_v2/figs/BorisEdits/fig1_scheme.png",
    "fig2": ROOT + "/docs/paper/draft_v2/figs/latest_review/na007b_3way_panel_polar.png",
    "fig3": ROOT + "/docs/paper/draft_v2/figs/BorisEdits/fig3_clustering.png",
    "fig4": ROOT + "/docs/paper/draft_v2/figs/latest_review/imc_param_maps_heat.png",
    "fig5": ROOT + "/docs/paper/draft_v2/figs/latest_review/imc_crystallization_schematic.png",
    "fig6": ROOT + "/docs/paper/draft_v2/figs/BorisEdits/fig6_gui_real.png",
    "si_vit": ROOT + "/docs/paper/draft_v2/figs/BorisEdits/fig_ablation_vit_vs_resnet.png",
    "si_depth": ROOT + "/docs/paper/draft_v2/figs/BorisEdits/fig_ablation_depth.png",
    "si_c1d": ROOT + "/docs/paper/draft_v2/figs/BorisEdits/fig_ablation_cluster1d.png",
    "si_box": ROOT + "/docs/paper/draft_v2/figs/latest_review/imc_class_boxplots_3col.png",
    "si_perframe": ROOT + "/docs/paper/draft_v2/figs/latest_review/imc_param_maps_heat_perframe.png",
    "si_perframe_bin": ROOT + "/docs/paper/draft_v2/figs/latest_review/imc_param_maps_heat_perframe_bin2.png",
    "si_si6_maps": ROOT + "/docs/paper/draft_v2/figs/latest_review/fig4_extra_SI6.png",
    "si_si6_box": ROOT + "/docs/paper/draft_v2/figs/latest_review/boxplots_extra_SI6.png",
    "si_si2_maps": ROOT + "/docs/paper/draft_v2/figs/latest_review/fig4_extra_SI2.png",
    "si_si2_box": ROOT + "/docs/paper/draft_v2/figs/latest_review/boxplots_extra_SI2.png",
    "si_si1_maps": ROOT + "/docs/paper/draft_v2/figs/latest_review/fig4_extra_SI1.png",
    "si_si1_box": ROOT + "/docs/paper/draft_v2/figs/latest_review/boxplots_extra_SI1.png",
    "si_naphi_dino": ROOT + "/docs/paper/draft_v2/figs/BorisEdits/fig2si_dino_classes.png",
    "si_naphi_nmf": ROOT + "/docs/paper/draft_v2/figs/BorisEdits/fig2si_nmf_classes.png",
    "si_naphi_iface": ROOT + "/docs/paper/draft_v2/figs/latest_review/na007b_class2v5v10.png",
    "si_nmf_extra": ROOT + "/docs/paper/draft_v2/figs/BorisEdits/nmf_hdbscan_fcm.png",
    "si_grainprov": ROOT + "/docs/paper/draft_v2/figs/latest_review/grain_locations_si.png",
    "si_carbon": ROOT + "/docs/paper/draft_v2/figs/latest_review/carbon_vs_imc_check.png",
    "si_alphafit": ROOT + "/docs/paper/draft_v2/figs/latest_review/si_precursor_alpha_fit.png",
    "si_gui": ROOT + "/docs/paper/draft_v2/figs/BorisEdits/fig6_gui_real.png",
    "fig4_scheme": ROOT + "/docs/paper/draft_v2/figs/latest_review/imc_scheme_only.png",
    "si_orderaxis": ROOT + "/docs/paper/draft_v2/figs/latest_review/imc_order_axis_example.png",
}

# AM Communication layout (4-figure cap): fold the mechanistic scheme into Figure 4
# and move the single-grain order-axis example to the SI. COMM=False -> 5-figure Article.
COMM = True
doc = Document()
st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(11)

# ===== Citations: AM style (superscript bracketed numbers, auto-numbered by order of appearance) =====
# In body text, mark a citation with {{key}} or several with {{key1;key2}}. Numbers are
# assigned the first time a key is seen, so the reference list comes out in appearance order.
_CITE_RE = re.compile(r"\{\{([^}]+)\}\}")
NUM = {}; _NEXT = [0]


def _numof(key):
    if key not in NUM:
        _NEXT[0] += 1; NUM[key] = _NEXT[0]
    return NUM[key]


def _compress(nums):
    nums = sorted(set(nums)); out = []; i = 0
    while i < len(nums):
        j = i
        while j + 1 < len(nums) and nums[j + 1] == nums[j] + 1:
            j += 1
        if j == i:
            out.append(str(nums[i]))
        elif j == i + 1:
            out.append(f"{nums[i]},{nums[j]}")
        else:
            out.append(f"{nums[i]}–{nums[j]}")
        i = j + 1
    return ",".join(out)


# ---- SI figure auto-numbering: {{S:key}} -> "Figure S<n>" by order of first reference ----
SNUM = {}; _SNEXT = [0]


def _snumof(key):
    if key not in SNUM:
        _SNEXT[0] += 1; SNUM[key] = _SNEXT[0]
    return SNUM[key]


def _fmt_si(nums):
    nums = sorted(set(nums))
    if len(nums) == 1:
        return f"Figure S{nums[0]}"
    if nums == list(range(nums[0], nums[-1] + 1)):
        return f"Figures S{nums[0]}–S{nums[-1]}"
    return "Figures " + ", ".join(f"S{n}" for n in nums)


def _emit_runs(p, text, size=None):
    """Add `text` to p, rendering {{key}} as a superscript AM citation and
    {{S:key}} (or {{S:k1;k2}}) as a normal 'Figure S<n>' cross-reference."""
    pos = 0
    for m in _CITE_RE.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            if size: r.font.size = size
        body = m.group(1)
        if body.startswith("S:"):
            nums = [_snumof(k.strip()) for k in body[2:].split(";")]
            r = p.add_run(_fmt_si(nums))
            if size: r.font.size = size
        else:
            nums = [_numof(k.strip()) for k in body.split(";")]
            r = p.add_run("[" + _compress(nums) + "]"); r.font.superscript = True
            if size: r.font.size = size
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        if size: r.font.size = size


def h(text, lvl=1):
    return doc.add_heading(text, level=lvl)


def para(text, justify=True):
    p = doc.add_paragraph()
    _emit_runs(p, text)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    return p


def figure(key, caption_bold, caption_rest):
    doc.add_picture(F[key], width=Inches(6.3))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    r = p.add_run(caption_bold); r.bold = True; r.font.size = Pt(9.5)
    _emit_runs(p, caption_rest, size=Pt(9.5))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(14)


def figure_multi(keys_widths, caption_bold, caption_rest):
    for key, w in keys_widths:
        doc.add_picture(F[key], width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    r = p.add_run(caption_bold); r.bold = True; r.font.size = Pt(9.5)
    _emit_runs(p, caption_rest, size=Pt(9.5))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(14)


def todo(text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.italic = True; r.font.color.rgb = RGBColor(0xB0, 0x40, 0x00)
    p.paragraph_format.space_after = Pt(12)


def expp(label, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(label + ": "); r.bold = True
    _emit_runs(p, text); p.paragraph_format.space_after = Pt(8)
    return p


# ===================== TITLE =====================
t = doc.add_heading("", level=0)
tr = t.add_run("Watching order emerge: a label-free self-supervised classifier for "
               "low-dose 4D-STEM of beam-sensitive organic crystals")
tr.bold = True
doc.add_paragraph("Daniel Khaykelson et al.").alignment = WD_ALIGN_PARAGRAPH.CENTER

# ===================== ABSTRACT =====================
h("Abstract", 1)
para(
"How crystals nucleate and grow from disordered organic solids is a central question for pharmaceuticals "
"and organic electronics, where the polymorph and the degree of order govern stability and performance. "
"Many such transformations are thought to be non-classical, proceeding through partly ordered precursors "
"rather than a sharp amorphous-to-crystalline jump, but resolving this at the nanoscale in beam-sensitive "
"films has remained a challenge. Four-dimensional scanning transmission electron microscopy (4D-STEM) "
"records a diffraction pattern at every probe position, but turning the resulting thousands of weak, "
"low-dose patterns into a map of structural order is the bottleneck, and has so far demanded "
"considerable manual effort and expertise. Here we address this with DINO4DSTEM, a label-free, "
"self-supervised classifier that sorts an entire low-dose 4D-STEM "
"dataset into structural domains automatically. On sodium poly(heptazine imide) it reproduces an "
"independent expert segmentation and resolves domains that automatic non-negative matrix factorization "
"misses. Applied to annealed indomethacin thin films, it shows that the crystalline kinetic α-phase "
"needles grow from a partly ordered, not amorphous, precursor: across the growth front the azimuthal "
"spottiness, two-dimensional Bragg excess, and radial peak-to-halo ratio rise continuously at fixed "
"d-spacing, identifying crystallization as the progressive azimuthal sharpening of a pre-existing "
"diffraction halo rather than the emergence of a new phase. Because the metastable α phase sets the "
"dissolution rate and bioavailability of the drug, resolving how and where it forms bears directly on "
"the control of pharmaceutical solids. Released as open software, the "
"method turns noisy, low-dose 4D-STEM datasets into reproducible nanoscale maps of order in "
"beam-sensitive organic materials.")

# ===================== INTRODUCTION =====================
h("1. Introduction", 1)
para(
"How order first emerges in an organic molecular crystal is a question of real practical weight: in "
"pharmaceuticals and organic electronics the polymorph and the degree of order set solubility, "
"stability, and charge transport.{{Bernstein2002}} That emergence is often not a simple jump from amorphous to "
"crystalline. In many soft and inorganic systems crystallization is non-classical, passing through "
"partly ordered precursors and intermediate states before the final phase appears.{{Du2024}} Resolving that "
"progression in a real organic film, at the nanoscale and grain by grain, would show directly how order "
"develops, yet it remains difficult to access: the same weak, non-covalent bonding that makes these "
"materials prone to polymorphism and defects also makes them acutely beam-sensitive, so the structural "
"information that would answer the question is both valuable and unusually hard to obtain. Indomethacin, "
"a poorly water-soluble pharmaceutical, is a canonical case: it is polymorphic, and which form grows, the "
"metastable α or the stable γ, sets its dissolution rate and hence its bioavailability, yet how the "
"crystalline phase first emerges from the amorphous solid has not been observed directly at the "
"nanoscale. It is the model material throughout this work.")
para(
"Four-dimensional scanning transmission electron microscopy (4D-STEM) is, in principle, the ideal probe "
"for this question.{{Ophus2019}} Fast direct-electron detectors, energy filters, and cryogenic stages now record a "
"full nanobeam diffraction pattern at every probe position, at doses low enough for beam-sensitive "
"organic films, producing thousands of patterns per scan. The limit is no longer the measurement but "
"the analysis: turning thousands of weak, low-dose patterns, some carrying clear scattering and many "
"close to the noise floor, into a coherent map of structural order is slow and leans heavily on the "
"judgement of an experienced operator.")
para(
"Several established and powerful methods already address parts of this problem, each suited to a "
"particular regime. Peak finding and automated orientation mapping{{Ophus2019}} excel when reflections are "
"sharp and a structure model is available, and are naturally less suited when the diffraction is diffuse "
"or only partially crystalline. Non-negative matrix factorization (NMF){{LeeSeung1999}} is flexible and "
"model-free, and in return asks the user to set the number of components and a subsequent clustering "
"step. Supervised segmentation, such as the Segment Anything Model (SAM),{{Kirillov2023}} is accurate and "
"intuitive, and benefits from human-guided sorting of the detected features for each new sample. What "
"these approaches share is a reliance on per-sample tuning and on microscopy expertise that is not "
"always available outside the field, which is the gap we set out to close, rather than a shortcoming of "
"any one method. Unsupervised representation learning, including autoencoders and variational "
"autoencoders, has also been applied to diffraction data; such models are typically not rotation "
"invariant, so grains of one phase imaged at different orientations are split into separate clusters, "
"which is precisely the difficulty our design is built to remove.")
para(
"Self-supervised representation learning offers a way out, by letting the data organize themselves "
"without labels or a preset number of classes,{{Caron2021}} and the idea maps cleanly onto crystallography. The "
"network is shown two versions of the same diffraction pattern that differ only by a transformation we "
"choose, and it is trained to give them the same representation; in learning to ignore that "
"transformation, it learns what is physically essential and discards what is not. We make the "
"transformation an in-plane rotation, so requiring a pattern and its rotated copy to be classified "
"identically is exactly the statement that a grain and the same grain turned to a different angle belong "
"together. The network therefore acquires orientation invariance directly, grouping grains of the same "
"phase and texture regardless of how they are oriented, which is the equivalence a crystallographer "
"would otherwise impose by hand. Models pre-trained on natural photographs, by contrast, transfer "
"poorly here, because a diffraction pattern carries far less, and very different, information than an "
"everyday image.")
para(
"In this work we implement these ideas as DINO4DSTEM, a self-supervised classifier built for low-dose "
"4D-STEM of beam-sensitive organic materials, adapting it to diffraction with a compact encoder and a "
"physics-based radial term that ties each class to its azimuthally averaged scattering (Results and "
"Experimental Section). We first benchmark it on NaPHI against a published expert "
"segmentation,{{Khaykelson2025}} then follow solid-state crystallization in annealed indomethacin films, where it "
"reveals that the α-needles grow from a partly ordered precursor by the progressive sharpening of a "
"diffraction halo that is already present. Because the classifier clusters the noisy frames before any "
"structural metric is computed, this transition becomes measurable where it is invisible to both "
"conventional imaging and frame-by-frame analysis. The accompanying software, including a graphical "
"interface and a natural-language assistant, is released openly (see Experimental Section and "
"github.com/DanielKhaykelson/dino-4dstem) so the analysis can be reproduced and reused.")

# ===================== 2. RESULTS AND DISCUSSION =====================
h("2. Results and Discussion", 1)
h("2.1. A self-supervised classifier adapted to diffraction", 2)
para(
"Building on the principle above, DINO4DSTEM groups diffraction patterns by similarity, the rotation "
"between the two training views supplying the orientation invariance, and two further adaptations make "
"it work on diffraction rather than on natural images. First, the encoder is kept deliberately small, "
"both because a diffraction pattern contains far less information than a photograph and because a fast, "
"light model is more likely to be adopted in practice than a slow, cumbersome one. Second, a "
"radial-profile term in the training objective ties each class to its azimuthally averaged scattering, "
"anchoring the classes to a physical quantity. In the layered samples we tested it also suppressed "
"leakage of classes across structural boundaries; we have not surveyed the full range of materials, so "
"we treat it as a useful option rather than a universal requirement, and report the ablation in the "
"Supporting Information. The choice of a compact convolutional encoder over a vision "
"transformer,{{Dosovitskiy2021}} its depth, and the radial-clustering loss are examined in the Supporting "
"Information ({{S:vit;depth;c1d}}). The same model settings were used for every dataset in this study, "
"with only routine per-sample pre-processing (centering, central-beam masking, and intensity scaling), "
"underscoring that the result does not rely on manual, per-sample tuning of the network; full settings "
"are given in the Experimental Section.")
figure("fig1",
"Figure 1. A label-free self-supervised classifier for 4D-STEM. ",
"A 4D-STEM scan records one nanobeam diffraction pattern at each probe position. Each pattern is shown "
"to the network as two augmented views, which a student encoder and a slowly updated teacher encoder "
"map into a shared representation; training drives the two views together, and a set of prototypes "
"groups the embeddings into classes. Taking the most likely prototype at each position gives the class "
"(domain) map. The augmentation that sets crystal orientation is the in-plane rotation, so the classes "
"are invariant to orientation. Training uses no labels and no preset number of classes; the objective "
"combines self-distillation with the one-dimensional radial clustering loss. The output panel is a "
"measured DINO4DSTEM class map (indomethacin). Encoder depth, teacher update, and loss weights are "
"given in the Experimental Section.")

# ===================== NaPHI =====================
h("2.2. Benchmarking against an expert segmentation (NaPHI)", 2)
para(
"We first tested DINO4DSTEM on NaPHI, where an independent reference is available. NaPHI is a layered, "
"two-dimensional organic photocatalyst whose flakes show diffuse, line-like scattering that arises from "
"internal buckling.{{Khaykelson2025}} These line domains make a useful benchmark: they are inaccessible to "
"peak-identification methods, yet pronounced enough to define a reliable reference. In earlier work "
"they were mapped with the Segment Anything Model (SAM),{{Khaykelson2025}} which gives accurate line-feature "
"parameters, such as the length and position of the line on the diffraction pattern, but is slow and "
"requires the user to sort the detected features. We therefore compared DINO4DSTEM with the published "
"SAM result and with polar non-negative matrix factorization followed by k-means,{{LeeSeung1999}} a standard "
"unsupervised approach for data that mix crystalline, amorphous, and diffuse-line scattering.")
para(
"All three methods recover the line domain, but they differ in fidelity and in what else they show "
"(Figure 2). The line domain is marked in blue throughout: as a SAM overlay on the virtual HAADF "
"image (Figure 2a), and as full class maps for NMF with k-means (Figure 2d) and DINO4DSTEM (Figure 2g). "
"The cluster maps reproduce the bright-field structure, separating vacuum, the NaPHI flake, and the "
"carbon support. Both unsupervised methods place the line domain in close agreement with the SAM "
"reference, overlapping it over about three-quarters of its area (spatial intersection-over-union 0.77 "
"for NMF and 0.74 for DINO4DSTEM). Because the line is a diffuse, "
"azimuthal feature rather than a sharp radial one, spatial overlap is the appropriate measure here; a "
"radial-profile correlation is dominated by the rings common to all NaPHI patterns and does not "
"distinguish the domains. Within this agreement, DINO4DSTEM separates the flake from the carbon support "
"more cleanly than NMF and resolves finer sub-domains inside the non-line flake.")
para(
"These sub-domains are not artefacts; they correspond to genuine differences in the diffraction. The "
"additional DINO4DSTEM classes track variations in specimen thickness and a measurement-induced "
"interface between the line region and the rest of the flake, each with a distinct average pattern "
"(Figure 2g and {{S:naphi_dino}}). Where this extra detail is not needed, the over-segmentation is "
"easily undone: the interface provides several merging protocols (combining classes a user considers "
"equivalent by spatial adjacency, diffraction similarity, or manual selection) without retraining. On "
"this benchmark, DINO4DSTEM is therefore as accurate as the supervised reference and more informative "
"than the standard unsupervised pipeline, on exactly the kind of low-dose, non-trivial organic "
"diffraction that motivates this work. The model, interface, and merging protocols are openly available "
"(github.com/DanielKhaykelson/dino-4dstem).")
figure("fig2",
"Figure 2. Validation against an independent expert segmentation (NaPHI). ",
"Rows: SAM (reference), rotation-invariant polar-theta-shift NMF, and DINO4DSTEM. Column 1 shows the "
"oriented line domain: SAM as a blue overlay on the virtual HAADF image (a), and NMF (d) and "
"DINO4DSTEM (g) as full class maps with the line cluster forced to the same blue and the remaining "
"classes in distinct colours. Columns 2 and 3 show the class-averaged diffraction of the line and "
"non-line regions. All three methods recover the line domain (spatial intersection-over-union against "
"SAM of 0.77 for NMF and 0.74 for DINO4DSTEM; because the line is a diffuse azimuthal feature, spatial "
"overlap rather than a radial-profile correlation is the meaningful measure). DINO4DSTEM additionally "
"separates flake from carbon and resolves finer, physically distinct sub-domains, namely thickness "
"variation and a measurement-induced interface, analysed per class in {{S:naphi_dino}}, with the full "
"set of NMF classes and their average diffraction in {{S:naphi_nmf}} and the line/rest interface domain "
"in {{S:naphi_iface}}.")

# ===================== IMC clustering (Fig 3) =====================
h("2.3. Mapping polycrystallinity in indomethacin thin films", 2)
para(
"We then applied the method to a partially crystallized indomethacin film. Solid-to-solid "
"crystallization of organic pharmaceuticals is a difficult target for diffraction: it requires a low "
"dose, not all of the reflections expected from the crystal structure are present in the early stages, "
"and the sample spans a range of orientations and degrees of crystallization.{{Du2024}} We grew 150 nm "
"indomethacin films by physical vapour deposition under conditions that give an amorphous film,{{Dawson2009}} then "
"annealed them at 70 °C, below the 155–162 °C melt, under nitrogen. The virtual HAADF images show "
"crystalline needles growing out of a matrix with much weaker structural definition (Figure 3, column "
"1), which makes the growth front a good place to study how order develops.")
para(
"Here the conventional real-space reference is not informative, which is precisely why an unsupervised "
"diffraction classifier is useful. The precursor matrix appears almost featureless in HAADF, so there "
"is no clear morphology against which to check a clustering, and no ground truth as there was for "
"NaPHI. We compared DINO4DSTEM with polar NMF, which requires a separate clustering step to turn its "
"loadings into a map, under several standard choices for that step (Figure 3). Two differences stand "
"out. First, the NMF map depends on the clustering step rather than being uniquely determined: applied "
"to the same loadings, k-means, agglomerative, and Gaussian-mixture clustering return visibly different "
"partitions, agreeing with one another only at an adjusted Rand index of 0.15 to 0.69{{Hubert1985}} "
"({{S:nmf_extra}} adds HDBSCAN and fuzzy-c-means, which differ further). This sensitivity is a natural "
"property of a flexible, model-free method, but in practice it means the user, rather than the data, "
"sets where the domain boundaries fall. DINO4DSTEM, using the same settings as for every other sample, "
"returns a single map with no such choice; its agreement with any individual NMF variant is "
"correspondingly modest (adjusted Rand index 0.05 to 0.27 across algorithms and the three regions), as "
"expected when comparing against a family of partitions rather than a single one.")
para(
"The result is a reproducible domain map of a sample whose structure is otherwise hard to access. "
"DINO4DSTEM reconstructs the field of view consistently across the three indomethacin regions, most "
"clearly where the crystalline needles give the strongest signal. Because it needs no per-sample "
"tuning and no visible morphological reference, the method reduces an ambiguous, operator-dependent "
"clustering task to a single automatic step, which is the basis for the physical analysis that "
"follows.")
figure("fig3",
"Figure 3. Clustering of the indomethacin fields of view. ",
"Rows are the three regions (interface, needles, and the magnified interface). Column 1 is the "
"virtual HAADF image with the 4D-STEM scan region; column 2 is the DINO4DSTEM class map; the remaining "
"columns are a standard polar NMF of the same patterns clustered by k-means, agglomerative, and "
"Gaussian-mixture algorithms. Colours are arbitrary cluster labels. DINO4DSTEM, with the same settings "
"used throughout, returns one consistent map per region, whereas NMF, with parameters optimized for "
"this sample, gives a different partition for each algorithm. Agreement between the NMF and DINO4DSTEM "
"maps is low and algorithm-dependent (adjusted Rand index 0.05 to 0.27), and the three NMF clusterings "
"agree even among themselves only at 0.15 to 0.69. Two further NMF clusterings (HDBSCAN and "
"fuzzy-c-means) and the per-algorithm comparison are in {{S:nmf_extra}}.")

# ===================== IMC crystallinity (Fig 4) =====================
h("2.4. A non-classical crystallization pathway in indomethacin", 2)
para(
"Sorting the single frames in this way places the indomethacin data on a single order axis. To read out "
"what the classifier has learned, we measured three rotation-invariant descriptors on the grain-average "
"diffraction of each class: the azimuthal spottiness (how strongly the ring intensity concentrates into "
"discrete spots), the two-dimensional Bragg excess B (integrated intensity above the smooth halo), and "
"the radial peak-to-halo ratio χ. Across the three co-located fields (interface, needles, and the "
"magnified interface; N = 49, 43, and 53 grains, respectively), all three descriptors rise monotonically "
"with the class label and the classes form tight, well-separated strata: the class label explains "
"η² = 0.54 to 0.89 of the spottiness variance, 0.42 to 0.87 of B, and 0.44 to 0.75 of χ (Figure 4; "
"per-class distributions in {{S:box}}). The class-median spottiness climbs continuously from about 0.2, "
"a near-uniform halo, to about 5, sharp discrete spots.")
para(
"This axis is the sharpening of one and the same reflection, not a sequence of distinct phases. In every "
"field the principal ring sits at d ≈ 4.5 to 4.8 Å, the α-indomethacin {102}/{112} spacing "
"(4.75 to 4.83 Å), and its position does not move along the order axis: the least-ordered classes "
"(azimuthal spottiness below about 0.5, a smooth amorphous-like halo) and the most-ordered (spottiness "
"of 4 to 5, discrete α Bragg spots) share the same ring radius to within the measurement scatter "
"(" + ("{{S:orderaxis}}" if COMM else "Figure 5") + ", "
"{{S:box}}). The least-ordered region is therefore not vacuum, carbon, or a structureless amorphous "
"solid: a diffraction-only test of the support confirms that bare amorphous carbon is featureless in our "
"window whereas this region shows the d ≈ 4.5 Å organic ring ({{S:carbon}}). It is a thick, "
"high-scattering area that already carries a weak but real α signal at the crystalline d-spacing, only "
"without azimuthal order; its azimuthally averaged profile indexes to the α reflections ({{S:alphafit}}). "
"Between it and the mature needles lie intermediate "
"classes that are unambiguously crystalline (non-zero Bragg excess, spottiness of order unity) yet do not "
"yet show the needle habit, so crystalline order is present before the morphology that usually announces "
"it. The needles thus grow from a partly ordered precursor by progressive azimuthal sharpening of a "
"pre-existing halo, the behavior reported for non-classical, multi-step crystallization across a range "
"of soft and inorganic systems.{{Du2024;Tsarfati2018;Zhang2022;Lin2023;Duan2023}} Because this precursor "
"looks non-crystalline in HAADF, the transition is hidden in conventional imaging and becomes measurable "
"only once the diffraction is sorted by order.")
para(
"This measurement is possible only because the clustering supplies the averaging that single frames "
"lack. On individual low-dose frames the descriptors are unreliable, because the per-frame signal is too "
"weak to rank the order robustly and can even invert it, and 2×2 detector binning does not restore the "
"class-level behavior "
"({{S:perframe}}). Clustering on the full two-dimensional pattern first, then measuring the descriptors on the "
"resulting class and grain averages, is therefore not a convenience but the step that makes the "
"measurement possible. With that averaging in hand the structural conclusion is firm: the least-ordered "
"region is a partly ordered, α-bearing precursor, and the needles arise from it by continuous azimuthal "
"sharpening at fixed d-spacing, not by nucleating a new phase. The data place this transition on a single "
"order axis and fix its structural signature unambiguously; what they do not yet resolve is its kinetics, "
"the rate and exact temporal sequence, which would require an in-situ anneal series.")
para(
"The same ordering ladder appears in three further fields of the same annealed film that were not part "
"of the analysis above. A separate needle field ({{S:si6}}; N = 66 grains, 8 classes, spottiness "
"η² = 0.68, class-median spottiness 0.20 to 3.21) reproduces it with the principal ring at d = 4.56 Å, "
"and two further regions reproduce it as well ({{S:si2}}, a thin film, N = 62 grains, η² = 0.71, ring "
"at 4.67 Å; {{S:si1}}, N = 68 grains, η² = 0.72, ring at 4.72 Å). Across all three the principal ring "
"stays at the α {102}/{112} spacing and the descriptors form the same monotonic strata, so the "
"fixed-d, azimuthal-sharpening signature is reproducible across fields of view and morphologies within "
"the film, not a feature of a single region.")
_FIG4_MAPS = (
"For each indomethacin region (rows), column 1 is the discrete DINO4DSTEM class map, with classes "
"coloured in order of azimuthal spottiness; columns 2 to 4 recolour the same classes by three "
"rotation-invariant descriptors measured per class, namely azimuthal spottiness, two-dimensional Bragg "
"excess, and radial peak-to-halo ratio, from dark (more amorphous) to bright (more crystalline). All "
"three vary as a coherent spatial gradient that peaks along the needles and is weakest at the "
"interface, identifying the discrete classes as steps along a single amorphous-to-α ordering axis "
"(N = 49, 43, and 53 grains for the three fields; η² = 0.54 to 0.89 for spottiness, 0.42 to 0.87 for "
"Bragg excess, and 0.44 to 0.75 for peak-to-halo; per-class distributions in {{S:box}}, the per-frame "
"and binned controls in {{S:perframe}}).")
if COMM:
    figure_multi([("fig4", 6.3), ("fig4_scheme", 6.0)],
    "Figure 4. The DINO4DSTEM classes trace a continuous crystallization axis, and the model it implies. ",
    "(Top) " + _FIG4_MAPS + " (Bottom) A schematic of the proposed mechanism: crystalline α needles grow "
    "from a partly ordered (not amorphous) precursor matrix, with the three imaged regions marked and an "
    "arrow indicating the gradual increase in structural order. Representative single-grain diffraction "
    "along this sequence is shown in {{S:orderaxis}}.")
else:
    figure("fig4",
    "Figure 4. The DINO4DSTEM classes follow a continuous crystallization-state axis. ", _FIG4_MAPS)
_scheme_ref = "Figure 4" if COMM else "Figure 5"
_example_clause = (" Representative single-grain patterns along the axis, and the provenance of those grains "
                   "on the scan, are given in {{S:orderaxis}} and {{S:grainprov}}." if COMM else "")
para(
"Together these measurements define a concrete, diffraction-based model for how indomethacin orders "
"(" + _scheme_ref + "). The α-needles grow from a partly ordered precursor by the progressive azimuthal sharpening "
"of a halo that is already present, at fixed d-spacing and without any new phase appearing at our "
"resolution. Crystallization here is therefore a gradual texturing of pre-existing scattering rather "
"than an abrupt amorphous-to-crystalline switch, a direct nanoscale realization of the non-classical, "
"precursor-mediated pathway that has been inferred more indirectly in other systems. That this pathway "
"can be read off a single annealed film, grain by grain, is the central result of this work, and it "
"follows directly from sorting the low-dose frames by order before measuring them." + _example_clause)
if not COMM:
    figure("fig5",
    "Figure 5. A proposed non-classical crystallization model for indomethacin. ",
    "(a) A schematic of the morphology, with crystalline α needles growing from a partly ordered (not "
    "amorphous) precursor matrix, the three imaged regions marked, and an arrow indicating the gradual "
    "increase in structural order. (b to d) Representative single-grain diffraction along the same "
    "sequence: (b) the partly ordered matrix (a strong principal-ring halo with only faint azimuthal "
    "order, spottiness ≈ 0.2), (c) the crystallization front (the ring condensing into sparse arcs, "
    "≈ 1.4), and (d) a mature needle (discrete α Bragg spots, ≈ 4.9), all at the same d-spacing. Ordering "
    "proceeds by azimuthal sharpening of a halo that is already present, not by the appearance of new "
    "phases. The provenance of the three grains on the scan is given in {{S:grainprov}} and the "
    "carbon-versus-indomethacin control in {{S:carbon}}.")

# ===================== Ease of use (demoted) =====================
h("2.5. Reproducibility and accessibility", 2)
para(
"A central aim of this work is to bridge the gap between electron microscopists and materials "
"scientists, putting this kind of analysis, and more, directly in the hands of non-specialists. Because "
"the network uses the same settings for every dataset, the only choices left to the user are physical "
"pre-processing steps, so the maps are reproducible across samples and operators rather than dependent "
"on per-sample tuning. The complete pipeline, including a graphical interface for multi-scale inspection "
"with gradient-based attribution (Grad-CAM){{Selvaraju2017}} and a natural-language assistant that runs the "
"workflow from plain-language requests, is released as open-source software (see the Experimental "
"Section, {{S:gui}}, and github.com/DanielKhaykelson/dino-4dstem).")

# ===================== Conclusion =====================
h("3. Conclusion", 1)
para(
"In annealed indomethacin, crystalline α-needles grow from a partly ordered precursor rather than from a "
"structureless amorphous solid: across the growth front the order rises continuously, from a "
"near-amorphous halo, through intermediate states that are already α but lack the needle habit, to "
"fully developed needles, all at the same d-spacing. The structural signature of this continuous, "
"precursor-mediated pathway, the progressive azimuthal sharpening of a pre-existing α ring, is fixed "
"unambiguously by the data; what remains for in-situ work is its kinetics, not whether it occurs. We "
"could read this pathway off a single annealed film only by sorting the low-dose frames by order before "
"measuring them, using a label-free, self-supervised 4D-STEM classifier (DINO4DSTEM); on the NaPHI "
"benchmark the same method reproduced an expert segmentation of the diffuse line domain and resolved "
"additional, physically real sub-domains.")
para(
"The same workflow should extend to other beam-sensitive molecular systems, including polymorph "
"screening, in-situ crystallization, and degradation studies, wherever weak, low-dose diffraction has "
"made large 4D-STEM datasets impractical to interpret by hand. Two properties make this realistic: the "
"model runs with the same settings across the datasets studied here, so results do not hinge on "
"per-sample tuning, and the full pipeline is openly available to groups outside the electron-microscopy "
"community. More generally, the work shows unsupervised clustering acting as the enabling measurement "
"step, recovering a physical signal that no single low-dose frame can provide.")

# ===================== Experimental Section =====================
h("4. Experimental Section", 1)
expp("Sample preparation",
"Indomethacin films (approximately 150 nm) were grown by physical vapour deposition at 0.2 Å s⁻¹ onto "
"lacy amorphous-carbon-coated gold TEM grids, forming an amorphous film, and then annealed ex situ at "
"70 °C for 60 min under nitrogen, above the glass transition (about 45 °C) and below the 155–162 °C "
"melt, to nucleate and grow crystalline α-needles. Polarized-light microscopy showed nucleation "
"starting at the lacy-carbon edges, consistent with strain induced on the film. The sodium "
"poly(heptazine imide) (NaPHI) flake and its reference SAM segmentation are from prior work.{{Khaykelson2025}}")
expp("4D-STEM acquisition",
"Energy-filtered nanobeam 4D-STEM was performed with a monochromated probe (16 eV energy-selecting "
"slit, beam current 3 to 4 pA, 20 µm condenser aperture, approximately 20 nm probe at −10 µm defocus) "
"at a camera length of 115 mm, giving a reciprocal-space calibration of 0.00185 Å⁻¹ per detector "
"pixel. Each dataset is a 128 × 128 probe-position scan (the central half of a 256 × 256 survey). The "
"interface and needle regions (SI3, SI4) were acquired at 7,100× magnification with a 44 nm step "
"(about 5.7 µm field) and 5 ms per pattern (dose about 50 e⁻ Å⁻²); the same needle/matrix "
"interface at higher magnification (SI5) was acquired at 20,000× with a 16 nm step (about 2.0 µm field) and 10 ms per pattern "
"(dose about 100 e⁻ Å⁻²). These doses are low enough to avoid beam-induced crystallization. "
"[To complete: microscope model and accelerating voltage, detector model and pixel pitch.]")
expp("Model and training",
"DINO4DSTEM uses self-distillation{{Caron2021}} with a student encoder and a slowly updated (exponential "
"moving-average) teacher encoder. The encoder is a ResNet-18{{He2016}} truncated to its first residual stage "
"(layer-1, two BasicBlocks), with deeper stages selectable up to four. The main augmentation is a "
"random in-plane rotation that enforces orientation invariance, applied together with light additional "
"augmentation (Gaussian blur and central-beam masking); a one-dimensional radial clustering loss "
"requires each pattern's azimuthally averaged profile to match its assigned class and to differ from "
"the others. An optional semi-supervised pair loss is available. The same hyperparameters were used "
"for every dataset in this study: 60 prototypes, 30 epochs, batch size 128, learning rate 3 × 10⁻⁴, "
"teacher momentum 0.97, teacher EMA from 0.99 to 0.999, and a polar input representation. Only "
"per-sample pre-processing was adjusted (centering, central-beam mask radius, and intensity scaling). "
"The choice of encoder, its depth, and the loss terms are examined in the Supporting Information ({{S:vit;depth;c1d}}).")
expp("Crystallinity descriptors",
"Three rotation-invariant descriptors were measured on class- and grain-average diffraction: the "
"azimuthal spottiness (90th percentile of the coefficient of variation of intensity around the "
"principal ring), the two-dimensional Bragg excess (integrated intensity above a smooth radial halo "
"baseline), and the radial peak-to-halo ratio. They are reported per class and per grain rather than "
"per frame, where shot noise dominates ({{S:perframe}}).")
expp("Validation of the precursor assignment",
"Three checks support interpreting the least-ordered on-sample region as a genuine, partly ordered "
"state. Its average diffraction indexes to the α-indomethacin structure, with the principal halo at "
"the {102}/{112} reflections near d ≈ 4.75 Å and a weaker feature at {103} (3.90 Å), at the same "
"d-spacings as the mature needles but azimuthally diffuse. The region is thick and high-scattering "
"(above-median scattered intensity) rather than thin or empty, so its single frames are shot-noise "
"limited. And the imaging dose (about 50 e⁻ Å⁻² for the interface and needle regions, 100 e⁻ Å⁻² for "
"the interface) is well below that associated with beam-induced crystallization, while radiation "
"damage would reduce order rather than create it.")
expp("Comparison methods and indexing",
"Non-negative matrix factorization{{LeeSeung1999}} was computed on a rotation-invariant polar representation and "
"clustered by k-means, agglomerative, Gaussian-mixture, HDBSCAN,{{McInnes2017}} and fuzzy-c-means; agreement with the "
"DINO4DSTEM maps is reported as the adjusted Rand index.{{Hubert1985}} The α-indomethacin reflections were assigned "
"from the α crystal structure{{IMCstruct}} by computed kinematical intensity (the {102}/{112} pair at "
"d ≈ 4.75–4.83 Å and {103} at 3.90 Å, with {022} and {003} at larger d).")
expp("Software, code, and data availability",
"The complete pipeline is released as open-source software at github.com/DanielKhaykelson/dino-4dstem. "
"It comprises the DINO4DSTEM model with pre-tuned parameters, a graphical interface, and a "
"natural-language assistant. The graphical interface lets a non-specialist load a dataset, run the "
"fixed-configuration model, and inspect the result at three scales, a single diffraction frame, a "
"grain, or a whole class, with gradient-based attribution (Grad-CAM){{Selvaraju2017}} indicating which part of "
"a pattern drove each class assignment, and offers several merging protocols to combine classes without "
"retraining. The natural-language assistant drives the same workflow, from loading and training through "
"clustering, inspection, and export, from plain-language requests, so no coding is required. Screenshots "
"and a walkthrough are in {{S:gui}}. [Data availability statement to complete.]")

# ===================== References (Advanced Materials style) =====================
# Numbered in order of first appearance (set by the {{key}} markers above). Journal
# abbreviations italic, year bold, volume italic, first page only.
h("References", 1)
REFDB = {
    "Bernstein2002": dict(kind="book", authors="J. Bernstein",
        title="Polymorphism in Molecular Crystals", publisher="Oxford University Press, Oxford, UK", year=2002),
    "Ophus2019": dict(kind="journal", authors="C. Ophus",
        journal="Microsc. Microanal.", year=2019, volume=25, pages=563),
    "LeeSeung1999": dict(kind="journal", authors="D. D. Lee, H. S. Seung",
        journal="Nature", year=1999, volume=401, pages=788),
    "Kirillov2023": dict(kind="proc", authors="A. Kirillov, E. Mintun, N. Ravi, H. Mao, C. Rolland, L. Gustafson, T. Xiao, S. Whitehead, A. C. Berg, W.-Y. Lo, P. Dollár, R. Girshick",
        venue="Proc. IEEE/CVF Int. Conf. Comput. Vis. (ICCV)", year=2023, pages=4015),
    "Caron2021": dict(kind="proc", authors="M. Caron, H. Touvron, I. Misra, H. Jégou, J. Mairal, P. Bojanowski, A. Joulin",
        venue="Proc. IEEE/CVF Int. Conf. Comput. Vis. (ICCV)", year=2021, pages=9650),
    "Khaykelson2025": dict(kind="raw",
        text="D. Khaykelson, et al., Nano Lett. 2025, DOI: [volume/page to complete]."),
    "Dosovitskiy2021": dict(kind="proc", authors="A. Dosovitskiy, L. Beyer, A. Kolesnikov, D. Weissenborn, X. Zhai, T. Unterthiner, M. Dehghani, M. Minderer, G. Heigold, S. Gelly, J. Uszkoreit, N. Houlsby",
        venue="Int. Conf. Learn. Represent. (ICLR)", year=2021, pages=None),
    "Du2024": dict(kind="journal", authors="J. S. Du, Y. Bae, J. J. De Yoreo",
        journal="Nat. Rev. Mater.", year=2024, volume=9, pages=923),
    "Dawson2009": dict(kind="journal", authors="K. J. Dawson, K. L. Kearns, L. Yu, W. Steffen, M. D. Ediger",
        journal="Proc. Natl. Acad. Sci. USA", year=2009, volume=106, pages=15165),
    "Hubert1985": dict(kind="journal", authors="L. Hubert, P. Arabie",
        journal="J. Classif.", year=1985, volume=2, pages=193),
    "Tsarfati2018": dict(kind="journal", authors="Y. Tsarfati, S. Rosenne, H. Weissman, L. J. W. Shimon, D. Gur, B. A. Palmer, B. Rybtchinski",
        journal="ACS Cent. Sci.", year=2018, volume=4, pages=1031),
    "Zhang2022": dict(kind="journal", authors="Z. Zhang, et al.",
        journal="Mater. Horiz.", year=2022, volume=9, pages=1670),
    "Lin2023": dict(kind="journal", authors="Y. Lin, F. Cheng, H. Wang, Y. Zhang, J. Fu, Y. Guo, J. Li, B. Ge",
        journal="Appl. Surf. Sci.", year=2023, volume=640, pages=158401),
    "Duan2023": dict(kind="journal", authors="T. Duan, Y. Shen, S. D. Imhoff, F. Yi, P. M. Voyles, J. H. Perepezko",
        journal="J. Chem. Phys.", year=2023, volume=158, pages="064504"),
    "Selvaraju2017": dict(kind="proc", authors="R. R. Selvaraju, M. Cogswell, A. Das, R. Vedantam, D. Parikh, D. Batra",
        venue="Proc. IEEE Int. Conf. Comput. Vis. (ICCV)", year=2017, pages=618),
    "He2016": dict(kind="proc", authors="K. He, X. Zhang, S. Ren, J. Sun",
        venue="Proc. IEEE Conf. Comput. Vis. Pattern Recognit. (CVPR)", year=2016, pages=770),
    "McInnes2017": dict(kind="journal", authors="L. McInnes, J. Healy, S. Astels",
        journal="J. Open Source Softw.", year=2017, volume=2, pages=205),
    "IMCstruct": dict(kind="raw",
        text="[Crystal-structure reference for α-indomethacin: to add]."),
}


def _add_reference(n, key):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25); p.paragraph_format.first_line_indent = Inches(-0.25)

    def rn(t, it=False, bd=False):
        r = p.add_run(t); r.font.size = Pt(9.5); r.italic = it; r.bold = bd

    ref = REFDB[key]; rn(f"[{n}] ")
    if ref["kind"] == "raw":
        rn(ref["text"]); return
    rn(ref["authors"] + ", ")
    if ref["kind"] == "journal":
        rn(ref["journal"] + " ", it=True); rn(str(ref["year"]), bd=True)
        rn(", "); rn(str(ref["volume"]), it=True); rn(f", {ref['pages']}.")
    elif ref["kind"] == "proc":
        rn("in " + ref["venue"] + ", ", it=True); rn(str(ref["year"]), bd=True)
        rn(f", p. {ref['pages']}." if ref.get("pages") else ".")
    elif ref["kind"] == "book":
        rn(ref["title"] + ", ", it=True); rn(ref["publisher"] + " "); rn(str(ref["year"]), bd=True); rn(".")


for key, n in sorted(NUM.items(), key=lambda kv: kv[1]):
    _add_reference(n, key)

# ===================== Supporting Information =====================
doc.add_page_break()
h("Supporting Information", 1)
para("This Supporting Information contains the architecture and loss ablations, the per-class "
"composition of the NaPHI segmentation, the clustering-dependence of NMF, the controls and "
"reproducibility tests behind the indomethacin crystallinity analysis, and the graphical interface. "
"Every DINO4DSTEM map uses the same training settings as the main results; only the component under "
"test, or the dataset, is varied.")


def sifigure(key, width, caption_bold, caption_rest):
    doc.add_picture(F[key], width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    r = p.add_run(caption_bold); r.bold = True; r.font.size = Pt(9.5)
    r2 = p.add_run(caption_rest); r2.font.size = Pt(9.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(14)


SIFIG = {
    "vit": ([("si_vit", 4.6)], "Encoder architecture.",
        "Class maps for NaPHI (Na007b) and indomethacin (SI5) using a vision-transformer encoder (a, c) "
        "versus the shallow convolutional encoder at matched depth (b, d). The transformer collapses to a "
        "few classes, while the convolutional encoder recovers the full domain structure. Each panel is "
        "annotated with the number of active classes and the effective number of classes."),
    "depth": ([("si_depth", 6.3)], "Encoder depth.",
        "Class maps as the convolutional encoder is deepened from its first residual stage (L1, two "
        "BasicBlocks) to four (L4), for NaPHI (a to d) and indomethacin (e to h). Shallow encoders (L1, "
        "L2) retain the domain structure; deeper encoders (L3, L4) collapse to fewer, coarser classes, "
        "which is why a shallow trunk is used throughout."),
    "c1d": ([("si_c1d", 5.6)], "One-dimensional radial clustering loss.",
        "Class maps with the loss off (a, c) and on (b, d) for the layered EuInAs sample and indomethacin "
        "(SI5). Each panel reports the active-class count, the effective number of classes, and the ratio "
        "of within-class to between-class feature distance (intra/inter; higher means better separated "
        "classes). For the layered EuInAs sample, where classes otherwise leak across the layer "
        "boundaries, the loss is needed to obtain a usable segmentation at all: it raises the separation "
        "(intra/inter 4.1 to 4.6) and removes a leaked class. For the non-layered indomethacin sample, "
        "which has no such leakage, it has little effect. The loss therefore matters specifically for "
        "leakage-prone samples, which can be the difference between a usable and an unusable result."),
    "naphi_dino": ([("si_naphi_dino", 6.5)], "NaPHI DINO4DSTEM classes and their meaning.",
        "The full set of DINO4DSTEM classes for the NaPHI flake (class map at left) with the average "
        "diffraction pattern of each class. The line domain, the thickness sub-domains, and the "
        "measurement-induced interface each have a distinct average pattern, confirming that the extra "
        "classes in Figure 2g are genuine structural differences rather than over-segmentation."),
    "naphi_nmf": ([("si_naphi_nmf", 6.5)], "NaPHI NMF classes.",
        "The corresponding NMF (K = 6, polar with θ-shift) class map and per-class average diffraction, "
        "clustered identically to Figure 2. NMF separates the broad regions but does not resolve the "
        "finer thickness and interface sub-domains that DINO4DSTEM recovers."),
    "naphi_iface": ([("si_naphi_iface", 6.3)], "The measurement-induced interface domain in NaPHI.",
        "The interface class and its neighbours at the boundary between the line domain and the rest of "
        "the flake, with average diffraction, radial profiles, and spatial location. About 69% of its "
        "pixels border the line region, identifying it as a thin transition rather than a separate phase."),
    "nmf_extra": ([("si_nmf_extra", 6.5)], "NMF under further clustering algorithms.",
        "The two NMF clusterings held out of Figure 3, HDBSCAN and fuzzy-c-means, for the three "
        "indomethacin regions, beside the DINO4DSTEM map. HDBSCAN collapses much of the field to noise "
        "and fuzzy-c-means gives yet another partition, so the NMF result depends strongly on the "
        "clustering algorithm, whereas the DINO4DSTEM map is fixed."),
    "box": ([("si_box", 6.5)], "Per-class distribution of the crystallization descriptors (indomethacin).",
        "Each box is one DINO4DSTEM class; dots are its individual grains; classes are ordered by median. "
        "Rows are the three co-located fields (interface, needles, magnified interface; N = 49, 43, and "
        "53 grains), columns the three rotation-invariant descriptors (azimuthal spottiness, "
        "two-dimensional Bragg excess B, radial peak-to-halo ratio χ). Measured per grain independently "
        "of the class label, the classes emerge as tight, separated, monotonic strata: the class label "
        "accounts for η² = 0.54 to 0.89 of the spottiness variance, 0.42 to 0.87 of B, and 0.44 to 0.75 "
        "of χ across the three fields. The class-median spottiness climbs continuously from about 0.2 (a "
        "near-uniform halo) to about 5 (discrete α Bragg spots)."),
    "carbon": ([("si_carbon", 6.3)], "Distinguishing the amorphous-carbon support from indomethacin.",
        "A diffraction-only discriminator. Amorphous carbon's main halo lies at d ≈ 2.1 Å, outside the "
        "recorded window, so bare support is nearly featureless in our frame, whereas a clear d ≈ 4.5 Å "
        "ring is an organic (indomethacin) signature. Radial profiles of the thickest and thinnest SI3 "
        "grains show that the least-ordered matrix region carries the 4.5 Å ring and is therefore "
        "indomethacin, not carbon."),
    "alphafit": ([("si_alphafit", 6.6)], "Indexing of the precursor and needle to α-indomethacin.",
        "Azimuthally averaged radial profiles of the least-ordered precursor and a mature needle on a "
        "reciprocal-space axis (q = 1/d; calibration 0.00185 Å⁻¹ per detector pixel), with the "
        "α-indomethacin reflections computed from the published crystal structure "
        "(alpha.cif; pymatgen, kinematical) overlaid as intensity-scaled sticks. The mature-needle peaks "
        "fall on the calculated α reflections across the whole range ({022} at 7.4 Å through {105} at "
        "3.1 Å), and the precursor halo sits under the same envelope, peaking at the principal {102}/{112} "
        "ring (4.75 to 4.83 Å). This confirms that the precursor is partly ordered α-indomethacin rather "
        "than a distinct phase, and that the reciprocal-space calibration is consistent with the "
        "structure."),
    "perframe": ([("si_perframe", 6.3), ("si_perframe_bin", 6.3)], "The descriptors are meaningful only after clustering.",
        "The same three descriptors measured on individual low-dose frames (top) and after 2×2 detector "
        "binning (bottom), rendered on the same maps as Figure 4. At the single-frame level shot noise "
        "dominates and the coherent spatial gradient of Figure 4 is absent; the apparent ordering is "
        "noisy and can locally invert. Binning does not recover the class-level behaviour, confirming "
        "that it is the clustering, not pixel-level averaging, that makes the order metrics measurable."),
    "si6": ([("si_si6_maps", 6.3), ("si_si6_box", 6.3)], "Separate indomethacin needle field (SI6).",
        "A separate needle field of the same annealed film, not used in the main analysis, processed "
        "identically. Top: the DINO4DSTEM class map and the three descriptor maps (as in Figure 4). "
        "Bottom: the per-class descriptor distributions (as in {{S:box}}). The same ordering ladder is "
        "reproduced (N = 66 grains, 8 classes, spottiness η² = 0.68, class-median spottiness 0.20 to "
        "3.21), with the principal ring at d = 4.56 Å, confirming that the fixed-d azimuthal-sharpening "
        "signature is not specific to one field of view."),
    "si2": ([("si_si2_maps", 6.3), ("si_si2_box", 6.3)], "Indomethacin thin-film region (SI2).",
        "A thin-film indomethacin region processed identically (N = 62 grains, 10 classes, spottiness "
        "η² = 0.71, class-median spottiness 0.79 to 5.31; principal ring d = 4.67 Å). The same "
        "monotonic order axis is recovered, showing that the result holds across sample geometries."),
    "si1": ([("si_si1_maps", 6.3), ("si_si1_box", 6.3)], "Second independent indomethacin region (SI1).",
        "A further indomethacin field (44 nm per pixel) processed identically (N = 68 grains, 11 classes, "
        "spottiness η² = 0.72, class-median spottiness 0.63 to 6.51; principal ring d = 4.72 Å, the "
        "closest of all fields to the nominal α {102}/{112} spacing). Top: the class map and three "
        "descriptor maps (as in Figure 4); bottom: the per-class distributions (as in {{S:box}}). The "
        "monotonic order axis and the fixed ring position are reproduced once more."),
    "grainprov": ([("si_grainprov", 6.3)], "Provenance of the representative diffraction patterns.",
        "For each of the three representative grains used to illustrate the order axis (the less-ordered "
        "matrix, the crystallization front, and the mature needle), the selected grain's footprint is "
        "marked on the scattered-intensity map of the scan, with its grain-average diffraction beneath, "
        "so each representative pattern is traceable to a location on the sample."),
    "gui": ([("si_gui", 6.5)], "Graphical interface and multi-scale inspection.",
        "The DINO4DSTEM graphical interface. The user loads a dataset, runs the fixed-configuration "
        "model, and inspects the result at the level of a single frame, a grain, or a whole class, with "
        "gradient-based attribution (Grad-CAM) indicating which part of a pattern drove each assignment. "
        "A natural-language assistant drives the same workflow from plain-language requests."),
    "orderaxis": ([("si_orderaxis", 6.6)], "The order axis forming, in single-grain diffraction.",
        "Representative grain-average diffraction along the crystallization sequence: (a) the partly "
        "ordered matrix (a strong principal-ring halo with only faint azimuthal order, spottiness ≈ 0.2), "
        "(b) the crystallization front (the ring condensing into sparse arcs, ≈ 1.4), and (c) a mature "
        "needle (discrete α Bragg spots, ≈ 4.9), all at the same d-spacing. The strip beneath each panel "
        "is that ring unrolled versus azimuth (flat = ring, peaks = discrete spots). Ordering proceeds by "
        "azimuthal sharpening of a pre-existing halo at fixed d-spacing, not by the appearance of a new "
        "phase. The provenance of the three grains on the scan is in {{S:grainprov}}."),
}


def _emit_sifig(n, key):
    images, title, body = SIFIG[key]
    for ik, w in images:
        doc.add_picture(F[ik], width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    r = p.add_run(f"Figure S{n}. {title} "); r.bold = True; r.font.size = Pt(9.5)
    _emit_runs(p, body, size=Pt(9.5))
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.paragraph_format.space_after = Pt(14)


for _k, _n in sorted(SNUM.items(), key=lambda kv: kv[1]):
    if _k in SIFIG:
        _emit_sifig(_n, _k)

out = ROOT + ("/docs/paper/draft_v2/draft_v2_comm_r4.docx" if COMM
              else "/docs/paper/draft_v2/draft_v2_full_r11.docx")
doc.save(out)
print("wrote", out)
